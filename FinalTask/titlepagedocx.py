import docx
from docx import Document
#from docx.shared import Inches
from docx.shared import Pt #импортируем размер шрифта
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH #импортируем равнение
import docxtpl
from docxtpl import DocxTemplate
import re

# создаем шаблон документа
doc = Document()

#устанавливаем поля
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    
#устанавливаем стиль текста
p_format = doc.styles['Normal'].paragraph_format
p_format.space_after = Pt(5)
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

#добавляем сам текст
firstblock = [doc.add_paragraph('Министерство науки и высшего образования Российской Федерации'),
              doc.add_paragraph('{{universityname}}'),
              doc.add_paragraph('{{facultyname}}'),
              doc.add_paragraph('{{cathedralname}}'),
              doc.add_paragraph(),
              doc.add_paragraph()]
for p in firstblock:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

secondblock = [doc.add_paragraph('ДОПУСТИТЬ К ЗАЩИТЕ В ГЭК'),
               doc.add_paragraph('Руководитель ООП'),
               doc.add_paragraph('{{achievements}}'),
               doc.add_paragraph('____________{{profname}}'),
               doc.add_paragraph('«_____»__________{{year}} г.'),
               doc.add_paragraph(),
               doc.add_paragraph(),
               doc.add_paragraph()]
for p in secondblock:
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА БАКАЛАВРА').bold = True

thirdblock = [doc.add_paragraph('{{nameofthesis}}'),
              doc.add_paragraph('по основной образовательной программе подготовки бакалавров\nнаправление подготовки {{numberofcourse}} – {{course}}'),
              doc.add_paragraph('{{fullstudentname}}'),
              doc.add_paragraph()]
for p in thirdblock:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

fourthblock = [doc.add_paragraph('Руководитель ВКР'),
               doc.add_paragraph ('{{leaderachievements}}'),
               doc.add_paragraph ('____________ {{leadername}}'),
               doc.add_paragraph ('«_____»__________{{year}} г.'),
               doc.add_paragraph(),
               doc.add_paragraph()]
for p in fourthblock:
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


fifthblock = [doc.add_paragraph ('Автор работы'),
              doc.add_paragraph ('студент группы № {{groupnumber}}'),
              doc.add_paragraph ('____________ {{studentname}}'),
              doc.add_paragraph()]
for p in fifthblock:
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    

doc.add_paragraph ('{{city}} – {{year}}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save ('titlepage.docx')

#подключаем получившийся шаблон
doc1 = DocxTemplate('titlepage.docx') 

universityname = str(input('Введите полное название вашего ВУЗа: ').strip())
while not re.match('^[A-Za-zА-Яа-яЁё\\s\'"()\/\\\.-]*$',universityname):
    print ('Проверьте правильность ввода данных.')
    universityname = str(input('Введите полное название вашего ВУЗа: ').strip())
facultyname = str(input('Введите полное название факультета: ').strip())
while not re.match('^[A-Za-zА-Яа-яЁё\\s\'"().-]*$',facultyname):
    print('Проверьте правильность ввода данных.')
    facultyname = str(input('Введите полное название факультета: ').strip())
cathedralname=str(input('Введите полное название кафедры: ').strip())
while not re.match('^[A-Za-zА-Яа-яЁё\\s\'"().-]*$',cathedralname):
    print('Проверьте правильность ввода данных')
    cathedralname=str(input('Введите полное название кафедры: ').strip())
profname=str(input('Введите инициалы и фамилию руководителя ООП, например И.И. Иванов: ').strip())
while not re.match('^[A-Za-zА-Яа-яёЁ\\s.-]*$', profname):
    print ('Пожалуйста, используйте только буквы алфавита, "." и "-"')
    profname=str(input('Введите инициалы и фамилию руководителя ООП, например И.И. Иванов: ').strip())
achievements=str(input('Введите должность/звание руководителя ООП (кратко): ').strip())
year = str(input('Введите год написания и защиты работы: ').strip())
while year.isdigit() == False:
    print ('Пожалуйста, вводите только цифры. ')
    year = str(input('Введите год написания и защиты работы: ').strip())
nameofthesis=str(input('Введите полное название вашей работы без кавычек: ').strip().upper())
numberofcourse=str(input('Введите номер направления ООП: ').strip())
course=str(input('Введите название ООП: ').strip())
fullstudentname = str(input('Введите ваше полное ФИО: ').strip())
while not re.match('^[A-Za-zА-Яа-яёЁ\\s.-]*$', fullstudentname):
    print ('Пожалуйста, используйте только буквы алфавита, "." и "-"')
    fullstudentname = str(input('Введите ваше полное ФИО: ').strip())
leadername=str(input('Введите инициалы и фамилию вашего научного руководителя, например И.И. Иванов: ').strip())
while not re.match('^[A-Za-zА-Яа-яёЁ\\s.-]*$', leadername):
    print ('Пожалуйста, используйте только буквы алфавита, "." и "-"')
    leadername=str(input('Введите инициалы и фамилию вашего научного руководителя, например И.И. Иванов: ').strip())
leaderachievements=str(input('Введите должность/звание вашего научного руководителя (кратко): ').strip())
studentname=str(input('Введите ваши инициалы и фамилию: ').strip())
while not re.match('^[A-Za-zА-Яа-яЁё\\s.-]*$',studentname):
    print ('Пожалуйста, используйте только буквы алфавита, "." и "-"')
    studentname=str(input('Введите ваши инициалы и фамилию: ').strip())
groupnumber=str(input('Введите номер группы: ').strip())
city=str(input('Введите ваш город: ').strip())
while not re.match('^[A-Za-zА-Яа-яЁё\\s\'"().-]*$',city):
    print ('Пожалуйста, проверьте правильность ввода данных')
    city=str(input('Введите ваш город: ').strip())
    
#переносим данные в словарь
context = {'universityname':universityname,'facultyname':facultyname,'cathedralname':cathedralname,
           'achievements':achievements,'profname':profname,'year':year,'nameofthesis':nameofthesis,
           'numberofcourse':numberofcourse,'course':course,'fullstudentname':fullstudentname,
           'leadername':leadername,'leaderachievements':leaderachievements,
           'studentname':studentname,'groupnumber':groupnumber,'city':city}
           
doc1.render(context)
doc1.save('titlepage.docx')
print('Смотрите файл "titlepage.docx" в той же папке, где запускали код :)')
